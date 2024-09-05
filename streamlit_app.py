import streamlit as st
import streamlit.components.v1 as components
from pyflowchart import *
from collections import OrderedDict
from docx import Document

st.set_page_config(layout="wide")
st.title("Procedure ➡ Flowchart Generator")
st.write(
    f"Generate a flowchart to visualize 'IF... GO TO' from Procedure"
)


def format_number(number_str):
    if number_str.count('.') == 1:
        number_str += '.0'
    if number_str.endswith('.'):
        number_str = number_str[:-1]
    return number_str[0:7]

upload = st.file_uploader("Select Procedure File", "docx")

doc = Document(upload)

class Step:
    def __init__(self, h1, h2, h3, text) -> None:
        self.node = None
        self.name = None
        self.h1 = h1
        self.h2 = h2
        self.h3 = h3
        self.text = text
        self.make_name()
        self.make_node()

    def __str__(self) -> str:
        return self.name

    def make_name(self) -> str:
        self.name = self.number + " " + self.text[:20]

    def accumulate(self, prev_node):
        self.name = f"{prev_node.name}\n{self.name}"
        prev_node.name = self.name  # Update the previous node's name

    def make_node(self) -> None:
        if self.is_conditional:
            self.node = ConditionNode(self)
        else:
            self.node = OperationNode(self)

    def goes_to(self) -> str:
        index = self.text.find("GO TO")
        if index != -1:
            t = self.text[index:]
            s = t.split()[3]
            return format_number(s)
        return None

    def links_to(self, prev_step):
        if prev_step.is_conditional:
        # if False:
            prev_step.node.connect_no(self.node)
        else:
            prev_step.node.connect(self.node)

    @property
    def is_conditional(self) -> bool:
        return self.text[:2] == "IF"

    @property
    def number(self) -> str:
        return f"{self.h1}.{self.h2}.{self.h3}"

heading_counts = {
    'Heading 1': 0,
    'Heading 2': 0,
    'Heading 3': 0,
}

h1 = 1
h2 = 1
h3 = 1

steps = OrderedDict()

for para in doc.paragraphs:
    if para.style.name in heading_counts:
        heading_counts[para.style.name] += 1
        if h1 != heading_counts['Heading 1']:
            heading_counts['Heading 2'] = 0
        if h2 != heading_counts['Heading 2']:
            heading_counts["Heading 3"] = 0
        h1 = heading_counts['Heading 1']
        h2 = heading_counts['Heading 2']
        h3 = heading_counts['Heading 3']
        if h1 > 4:
            step = Step(h1, h2, h3, para.text)
            steps[step.number] = step

index = list(steps)

for i in range(1, len(index)-1):
    current_step = steps[index[i]]
    previous_step = steps[index[i-1]]

    current_step.links_to(previous_step)
    if current_step.is_conditional and current_step.goes_to():
        if current_step.goes_to():
            if current_step.goes_to() in steps.keys():
                current_step.node.connect_yes(steps[current_step.goes_to()].node)


bg = StartNode('Procedure')
if upload:
    bg.connect(steps[index[0]].node)

fc = Flowchart(bg)

html_template = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0, initial-scale=0.8, minimum-scale=0.8">
    <title>{doc.paragraphs[0].text}</title>
    <link rel="stylesheet" href="https://ajax.googleapis.com/ajax/libs/jqueryui/1.12.1/themes/smoothness/jquery-ui.css">
    <script src="https://ajax.googleapis.com/ajax/libs/jquery/3.3.1/jquery.min.js"></script>
    <script src="https://ajax.googleapis.com/ajax/libs/jqueryui/1.12.1/jquery-ui.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/raphael/2.2.7/raphael.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/flowchart/1.8.0/flowchart.min.js"></script>
    <style>
        body {{
            font-family: Arial, sans-serif;
            padding: 20px;
            background-color: #f4f4f9;
            transform: scale(0.8); /* Scales down content */
            transform-origin: 0 0; /* Keeps scaling anchored to top-left */
        }}
        #diagram {{
            width: 100%;
            height: 500px;
            border: 1px solid #ccc;
        }}
    </style>
</head>
<body>

    <h1>{upload.name}</h1>
    <div id="diagram"></div>

    <script type="text/javascript">
        $(document).ready(function () {{
            var flowchartCode = `{fc.flowchart()}`;
            var diagram = flowchart.parse(flowchartCode);
            diagram.drawSVG('diagram');
        }});
    </script>

</body>
</html>
"""

components.html(html_template, height=500, scrolling=True, width=None)